"""
Document processing service for RAG QA Chatbot Application
"""
import uuid
import io
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple, Union
import PyPDF2
from docx import Document as DocxDocument

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

from ..config import config
from ..models import Document, DocumentChunk, DocumentType
from ..utils import (
    app_logger, validate_file_type, get_file_hash,
    sanitize_filename, format_file_size, measure_execution_time,
    retry_on_exception
)


class DocumentProcessor:
    """Document processing service"""

    def __init__(self):
        self.logger = app_logger
        self.supported_extensions = config.document.supported_extensions
        self.max_file_size = config.document.max_file_size_mb * \
            1024 * 1024  # Convert to bytes

    @measure_execution_time
    def process_uploaded_files(self, uploaded_files) -> List[Document]:
        """
        Process multiple uploaded files

        Args:
            uploaded_files: List of uploaded file objects

        Returns:
            List of processed Document objects
        """
        documents = []

        for uploaded_file in uploaded_files:
            self.logger.debug(
                f"Starting processing for file: {uploaded_file.name}")
            try:
                document = self.process_single_file(uploaded_file)
                if document:
                    documents.append(document)
                    self.logger.info(
                        f"Successfully processed file: {document.name}")
                else:
                    self.logger.warning(
                        f"File {uploaded_file.name} was not processed successfully.")
            except Exception as e:
                self.logger.error(
                    f"Failed to process file {uploaded_file.name}: {str(e)}")
                continue

        self.logger.info(
            f"Processed {len(documents)} out of {len(uploaded_files)} files")
        return documents

    @retry_on_exception(max_retries=3)
    def process_single_file(self, uploaded_file) -> Optional[Document]:
        """
        Process a single uploaded file

        Args:
            uploaded_file: Uploaded file object

        Returns:
            Document object or None if processing failed
        """
        self.logger.debug(f"Validating file: {uploaded_file.name}")
        # Validate file
        if not self._validate_file(uploaded_file):
            self.logger.warning(
                f"Validation failed for file: {uploaded_file.name}")
            return None

        # Read file content
        file_content = uploaded_file.read()
        uploaded_file.seek(0)  # Reset file pointer

        # Special handling for DOCX files - convert to PDF first
        file_type = self._get_file_type(uploaded_file.name)
        if file_type == DocumentType.DOCX and DOCX2PDF_AVAILABLE:
            try:
                self.logger.debug(
                    f"Converting DOCX to PDF: {uploaded_file.name}")
                file_content, file_type = self._convert_docx_to_pdf(
                    file_content, uploaded_file.name)
                self.logger.info(
                    f"Converted {uploaded_file.name} from DOCX to PDF for better page handling")
            except Exception as e:
                self.logger.warning(
                    f"Failed to convert DOCX to PDF, processing as DOCX: {str(e)}")

        # Extract text based on file type
        self.logger.debug(f"Extracting text from file: {uploaded_file.name}")
        text_content = self._extract_text(
            file_content, file_type, uploaded_file.name)

        if not text_content.strip():
            self.logger.warning(
                f"No text content extracted from {uploaded_file.name}")
            return None

        # Create document object
        original_file_type = self._get_file_type(uploaded_file.name)
        document = Document(
            id=str(uuid.uuid4()),
            name=sanitize_filename(uploaded_file.name),
            content=file_content,  # Store raw bytes (might be converted PDF)
            file_type=file_type,   # This might be PDF if DOCX was converted
            file_size=len(file_content),
            file_hash=get_file_hash(file_content),
            metadata={
                'original_name': uploaded_file.name,
                'original_file_type': original_file_type.value,  # Store original type
                'converted_from_docx': original_file_type == DocumentType.DOCX and file_type == DocumentType.PDF,
                'character_count': len(text_content),
                'word_count': len(text_content.split()),
                'text_content': text_content  # Store extracted text in metadata
            }
        )

        self.logger.debug(f"Created document: {document.name} "
                          f"({format_file_size(document.file_size)})")

        return document

    def _validate_file(self, uploaded_file) -> bool:
        """Validate uploaded file"""
        # Check file type
        if not validate_file_type(uploaded_file.name, self.supported_extensions):
            self.logger.warning(f"Unsupported file type: {uploaded_file.name}")
            return False

        # Check file size
        file_size = uploaded_file.size if hasattr(
            uploaded_file, 'size') else len(uploaded_file.getvalue())
        if file_size > self.max_file_size:
            self.logger.warning(f"File too large: {uploaded_file.name} "
                                f"({format_file_size(file_size)})")
            return False

        # Check if file is empty
        if file_size == 0:
            self.logger.warning(f"Empty file: {uploaded_file.name}")
            return False

        return True

    def _get_file_type(self, filename: str) -> DocumentType:
        """Get document type from filename"""
        suffix = Path(filename).suffix.lower()

        if suffix == '.pdf':
            return DocumentType.PDF
        elif suffix == '.docx':
            return DocumentType.DOCX
        elif suffix == '.txt':
            return DocumentType.TXT
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _convert_docx_to_pdf(self, file_content: bytes, filename: str) -> Tuple[bytes, DocumentType]:
        """
        Convert DOCX file content to PDF using simple docx2pdf conversion

        Args:
            file_content: DOCX file content as bytes
            filename: Original filename for logging

        Returns:
            Tuple of (pdf_content_bytes, PDF_DocumentType)
        """
        if not DOCX2PDF_AVAILABLE:
            raise ImportError("docx2pdf package not available")

        import pythoncom
        pythoncom.CoInitialize()  # Ensure COM is initialized

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir_path = Path(temp_dir)

            # Save DOCX to temporary file
            docx_path = temp_dir_path / filename
            with open(docx_path, 'wb') as f:
                f.write(file_content)

            # Convert DOCX to PDF using simple method
            pdf_path = temp_dir_path / f"{Path(filename).stem}.pdf"

            try:
                # Simple conversion: give docx file directly to convert function
                convert(str(docx_path), str(pdf_path))

                # Read converted PDF content
                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()

                self.logger.debug(
                    f"Successfully converted {filename} to PDF ({len(pdf_content)} bytes)")
                return pdf_content, DocumentType.PDF

            except Exception as e:
                self.logger.error(
                    f"docx2pdf conversion failed for {filename}: {str(e)}")
                raise

    def _extract_text(self, file_content: bytes, file_type: DocumentType, filename: str) -> str:
        """Extract text from file content based on file type"""
        try:
            if file_type == DocumentType.PDF:
                return self._extract_pdf_text(file_content)
            elif file_type == DocumentType.DOCX:
                return self._extract_docx_text(file_content)
            elif file_type == DocumentType.TXT:
                return self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            self.logger.error(
                f"Failed to extract text from {filename}: {str(e)}")
            raise

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF content"""
        import io

        text = ""
        pdf_file = io.BytesIO(file_content)

        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
                except Exception as e:
                    self.logger.warning(
                        f"Failed to extract text from page {page_num}: {str(e)}")
                    continue

        except Exception as e:
            self.logger.error(f"Failed to read PDF: {str(e)}")
            raise

        return text.strip()

    def _extract_pdf_text_by_pages(self, file_content: bytes) -> List[Tuple[str, int]]:
        """
        Extract text from PDF content page by page

        Returns:
            List of (page_text, page_number) tuples
        """
        import io

        pages_content = []
        pdf_file = io.BytesIO(file_content)

        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # 1-based page numbers
                        pages_content.append((page_text.strip(), page_num + 1))
                except Exception as e:
                    self.logger.warning(
                        f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue

        except Exception as e:
            self.logger.error(f"Failed to read PDF: {str(e)}")
            raise

        return pages_content

    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX content"""
        import io

        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            return text.strip()

        except Exception as e:
            self.logger.error(f"Failed to read DOCX: {str(e)}")
            raise

    def _extract_docx_text_by_pages(self, file_content: bytes) -> List[Tuple[str, int]]:
        """
        Extract text from DOCX content by pages (approximated)

        Note: DOCX doesn't have explicit page breaks, so we approximate
        by splitting content into logical sections

        Returns:
            List of (page_text, page_number) tuples
        """
        import io

        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)

            pages_content = []
            current_page_text = ""
            page_number = 1
            char_count = 0
            # Approximate characters per page (adjust as needed)
            chars_per_page = 2000

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_text = paragraph.text + "\n"
                    char_count += len(paragraph_text)
                    current_page_text += paragraph_text

                    # If we've reached approximate page size, start new page
                    if char_count >= chars_per_page:
                        if current_page_text.strip():
                            pages_content.append(
                                (current_page_text.strip(), page_number))
                        page_number += 1
                        current_page_text = ""
                        char_count = 0

            # Add remaining text as the last page
            if current_page_text.strip():
                pages_content.append((current_page_text.strip(), page_number))

            return pages_content

        except Exception as e:
            self.logger.error(f"Failed to read DOCX: {str(e)}")
            raise

    def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace')

        except Exception as e:
            self.logger.error(f"Failed to read TXT: {str(e)}")
            raise

    def _extract_txt_text_by_pages(self, file_content: bytes) -> List[Tuple[str, int]]:
        """
        Extract text from TXT content by pages (approximated)

        Returns:
            List of (page_text, page_number) tuples
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text_content = None

            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            if text_content is None:
                text_content = file_content.decode('utf-8', errors='replace')

            # Split into pages based on line count or character count
            lines = text_content.split('\n')
            pages_content = []
            current_page_lines = []
            page_number = 1
            lines_per_page = 50  # Approximate lines per page

            for line in lines:
                current_page_lines.append(line)

                if len(current_page_lines) >= lines_per_page:
                    page_text = '\n'.join(current_page_lines).strip()
                    if page_text:
                        pages_content.append((page_text, page_number))
                    page_number += 1
                    current_page_lines = []

            # Add remaining lines as the last page
            if current_page_lines:
                page_text = '\n'.join(current_page_lines).strip()
                if page_text:
                    pages_content.append((page_text, page_number))

            return pages_content

        except Exception as e:
            self.logger.error(f"Failed to read TXT: {str(e)}")
            raise

    @measure_execution_time
    def create_chunks(self, document: Document, model_provider: str, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None) -> List[DocumentChunk]:
        """
        Split document content into chunks with page information

        Args:
            document: Document to chunk
            model_provider: Model provider to determine chunk strategy
            chunk_size: Optional custom chunk size (overrides default)
            chunk_overlap: Optional custom chunk overlap (overrides default)

        Returns:
            List of DocumentChunk objects with page information
        """
        from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter

        try:
            # Use custom parameters if provided, otherwise use config defaults
            if model_provider == "OpenAI":
                default_chunk_size = config.document.openai_chunk_size
                default_chunk_overlap = config.document.chunk_overlap
            elif model_provider == "Google AI":
                default_chunk_size = config.document.google_chunk_size
                default_chunk_overlap = config.document.google_chunk_overlap
            else:
                default_chunk_size = config.document.chunk_size
                default_chunk_overlap = config.document.chunk_overlap

            # Use provided parameters or fall back to defaults
            final_chunk_size = chunk_size if chunk_size is not None else default_chunk_size
            final_chunk_overlap = chunk_overlap if chunk_overlap is not None else default_chunk_overlap

            # Get page-based content based on document type
            try:
                if document.file_type == DocumentType.PDF:
                    pages_content = self._extract_pdf_text_by_pages(
                        document.content)
                elif document.file_type == DocumentType.DOCX:
                    pages_content = self._extract_docx_text_by_pages(
                        document.content)
                elif document.file_type == DocumentType.TXT:
                    pages_content = self._extract_txt_text_by_pages(
                        document.content)
                else:
                    # Fallback to regular text extraction
                    pages_content = [
                        (document.content.decode('utf-8', errors='replace'), 1)]
            except Exception as e:
                self.logger.warning(
                    f"Failed to extract page-based content, falling back to regular extraction: {str(e)}")
                # Fallback to regular text extraction
                if isinstance(document.content, bytes):
                    text_content = document.content.decode(
                        'utf-8', errors='replace')
                else:
                    text_content = str(document.content)
                pages_content = [(text_content, 1)]

            # Configure text splitter based on model provider
            if model_provider == "OpenAI":
                text_splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size=final_chunk_size,
                    chunk_overlap=final_chunk_overlap,
                    length_function=len
                )
            elif model_provider == "Google AI":
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=final_chunk_size,
                    chunk_overlap=final_chunk_overlap
                )
            else:
                # Default splitter
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=final_chunk_size,
                    chunk_overlap=final_chunk_overlap
                )

            # Create chunks with page information
            all_chunks = []
            chunk_index = 0

            for page_text, page_number in pages_content:
                if not page_text.strip():
                    continue

                # Split page content into chunks
                page_chunks = text_splitter.split_text(page_text)

                for chunk_text in page_chunks:
                    if not chunk_text.strip():
                        continue

                    # Create chunk with page information in metadata
                    chunk = DocumentChunk(
                        id=str(uuid.uuid4()),
                        document_id=document.id,
                        content=chunk_text,
                        chunk_index=chunk_index,
                        start_char=0,  # Would need more complex calculation for exact positions
                        end_char=len(chunk_text),
                        metadata={
                            'page': page_number,  # Real page number from document
                            'document_name': document.name,
                            'file_type': document.file_type.value,
                            'chunk_size': len(chunk_text),
                            'model_provider': model_provider
                        }
                    )
                    all_chunks.append(chunk)
                    chunk_index += 1

            self.logger.info(f"Created {len(all_chunks)} chunks from {len(pages_content)} pages "
                             f"for document: {document.name}")
            return all_chunks

        except Exception as e:
            self.logger.error(
                f"Failed to create chunks for document {document.name}: {str(e)}")
            raise


__all__ = ["DocumentProcessor"]
